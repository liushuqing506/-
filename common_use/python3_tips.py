pycharm 往前缩进 Shift + Tab
dpkg --list 查看已经安装的软件

*************定义有序字典，值为列表
from collections import OrderedDict
D_lines_dict = OrderedDict()
D_lines_dict.setdefault(k, []).append(str or list)
or
import collections
read_alignments = collections.defaultdict(list)
read_alignments['ok'].append(['12','34'])

*************字典定义多个键
class AutoVivification(dict):
    "Implementation of perl's autovivification feature."
    def __getitem__(self, item):
        try:
            return dict.__getitem__(self, item)
        except KeyError:
            value = self[item] = type(self)()
            return value

info=AutoVivification()
info['A']['B']['C']['D'] = 'OK'

**************多键字典使用
domain_id_dict[domain][speice_id] = speice
for i,j in domain_id_dict.items():
    for m,n in j.items():

*************同时遍历两个列表或者字典
a = {'A':'1','C':'1'}
b = {'B':'2','D':'2'}

for i,k in zip(a.keys(),b.keys()):
    print(i+k)
*************三元表达式
j = [j for j in species_ZE_dict.keys() if j in i.strip()]
or
age = '' if jsonDict['age'] == None else str(jsonDict['age'])

*************Word表格处理<https://zhuanlan.zhihu.com/p/88151371>
import time
import datetime
import sys
import re
import docx
from docx import *
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm

def set_run(run, font_size=12, bold = False, R=0,G=0,B=0, name='等线',italic=False):
    '''
    设置run对象
    :param run:
    :param font_size: 字体大小
    :param bold: 是否加粗
    :param color: 字体颜色
    :param name: 字体名
    :return:
    '''
    run.font.size = Pt(font_size) #默认小四
    run.font.bold = bold  #默认不加粗
    run.font.color.rgb = RGBColor(R,G,B) #默认黑色
    run.font.name = name
    run.font.italic = italic #默认不斜体
    # 设置字体必须要下面2步
    s = run._element
    s.rPr.rFonts.set(qn('w:eastAsia'), name)


docfile = "demo.docx"

doc = Document(docfile)
doc.tables[2].cell(0,0).paragraphs[0].text #查看内容
len(doc.tables) #整个word文档表格总数

table = doc.tables[2]
len(table.rows), len(table.columns) #tables[2]表格的行数和列数
table.add_row() #底部增加一行
table.add_column(Cm(3))#右端增加一列（Cm(3)宽度）
run = doc.tables[0].cell(0,1).paragraphs[0].add_run('RICU') #添加对应位置数据
set_run(run)


#doc.tables[0].cell(1,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 居中
doc.save(FAH_num+"-"+name+"-化疗药物多态性"+"-"+"".join(report_date.split("-"))+".docx")

********************列表去重，但是顺序不变
list1 = [0, 3, 2, 3, 1, 0, 9, 8, 9, 7]
list2 = list(set(list1))
print(list2)        # [0, 1, 2, 3, 7, 8, 9]
list2.sort(key = list1.index)
print(list2)        # [0, 3, 2, 1, 9, 8, 7]


********************按字典的值排序
import operator
Bacteria_number_Genus_dict = {'A':1,'C':1,'B':2,'D':1,'E':3,}
print(sorted(Bacteria_number_Genus_dict.items(),key=operator.itemgetter(1), reverse=True))  #值<itemgetter(1)>，降序
返回[('E', 3), ('B', 2), ('A', 1), ('C', 1), ('D', 1)]
********按字典的键排序
print(sorted(Bacteria_number_Genus_dict.items(),key=operator.itemgetter(0),))  #键<itemgetter(1)>，升序
以上返回的都是列表，如果需要返回字典如下：
dict(sorted(Bacteria_number_Genus_dict.items(),key=operator.itemgetter(0),))

*********************列表去交集，并集，差集
交集 list(set(genus_abundance_list).intersection(set(species_abundance_list)))
差集genus_abundance_list   list(set(genus_abundance_list).difference(set(species_abundance_list)))
差集species_abundance_list   list(set(species_abundance_list).difference(set(genus_abundance_list)))
并集 list(set(a).union(b,c))
备注：https://www.cnblogs.com/yezishen/p/11739619.html

*********************同时开启两个文件
with open(abundance_file, "r",encoding='UTF-8') as f2,open(abundance_file, "w",encoding='UTF-8') as f3:

*********************python同时执行两个函数
import threading
import time
def minimap2(fastq_file,minimap2_q):
    ...
def kraken(fastq_file):
    ...
    
threads = []
threads.append(threading.Thread(target=minimap2, args=(args.fastq, args.q)))
threads.append(threading.Thread(target=kraken, args=(args.fastq,))) #参数必须使用（），如果有一个参数（args01，）
for t in threads:
    t.start()
for t in threads:
    t.join()   #join()方法 是等待线程结束，才执行后面的，不然线程直接进入后端之后就立即执行后面的

***********************统计起始时间
startTime = time.time()
...
endTime = time.time()
print(endTime-startTime)
--------------------------------
import time
start_time = time.perf_counter()
time.sleep(5) #中间过程。。。
stop_time = time.perf_counter()
cost = stop_time - start_time
with open('cost_{0}s'.format(cost),'w') as fw:
    pass

***********************统计路径下所有文件
file_list = []
for parent, dirnames, filenames in os.walk(path,  followlinks=True):
    for filename in filenames:
        file_path = os.path.join(parent, filename)
        file_list.append(file_path)
        print('文件完整路径：%s\n' % file_path)

************************保留位数
四舍五入
a = 12.345
round(a, 2)
12.35

'{:.2%}'.format(12.34532) #'1234.53%' 百分比并且小数点保留两位

*************************py3读取excle数据
import xlrd
worksheet = xlrd.open_workbook('XXXX.xlsx')   #打开excel文件
sheet_names= worksheet.sheet_names()    #获取excel中所有工作表名
sheet2 = worksheet.sheet_by_name('Sheet2')    #根据Sheet名获取数据
rows=sheet2.nrows  #Sheet2中行数
ncols = sheet1.ncols  #Sheet2中列数
rows = sheet2.row_values(3)   #表示获取Sheet2中第4行数据
cols10 = sheet2.col_values(9)   #表示获取Sheet2中第10列数据（数据保存为list）
cell_A1 = sheet1.cell(0,0).value #单元格内容

xlrd.xldate_as_tuple(sheet1.cell(i,j).value,worksheet.datemode) #读取excel中日期类型的数据（2020/7/2），返回（2020,7,2,0，0）

*************************把当前所在路径添加到环境变量中
临时 ： export PATH=$PATH:$(pwd)  
永久 ： 
vi ~/.bashrc   export PATH=$PATH:/home/liushuqing/miniconda3/envs/minimap2/bin/
. ~/.bashrc
  
************************* 特定目录下面的文件
file_list = os.listdir('./B_fastq/')

*************************文件是否存在？？
os.path.exists(coverage_png_file)

*************************word添加固定章节
设置章节,增加/减少/修改章节位置

***************************word模板的分页符
可以设置分页符来单独分页

*************************写入csv
import csv
with open('%s_sorted_abundance.csv'%classification,'w',newline='',encoding='utf-8-sig') as csvFile:
    writer = csv.writer(csvFile)
    writer.writerow(("Species_name", "S_Ch", "reads_number", "abundance<%>", "Genus_name", "G_Ch", "reads_number",\
                     "abundance<%>"))

*************************统计文件行数
file_lines = len(open('../test.fastq','r').readlines())

*************************异常捕获
blast_path = r'H:\python3\project\test\0821\B_blast_1'
try:
    file_list = os.listdir(blast_path)
except IOError as msg:
    print('异常：{0}'.format(msg))  
else:
    print('读取文件夹正常')
多种异常类型详情：https://www.runoob.com/python/python-exceptions.html

*************************列表按指定个数拆分为若干子列表
def list_of_groups(init_list, childern_list_len):
    '''
    init_list为初始化的列表，childern_list_len初始化列表中的几个数据组成一个小列表
    :param init_list:
    :param childern_list_len:
    :return:
    '''
    list_of_group = zip(*(iter(init_list),) *childern_list_len)
    end_list = [list(i) for i in list_of_group]
    count = len(init_list) % childern_list_len
    end_list.append(init_list[-count:]) if count !=0 else end_list
    return end_list

l = [i for i in range(15)]
print(list_of_groups(l,2)) #[[0, 1], [2, 3], [4, 5], [6, 7], [8, 9], [10, 11], [12, 13], [14]]

***************************查看外置硬盘挂载信息
fdisk -l 
df -hT
lsblk  -f
***************************查看内存大小
less /proc/meminfo 
MemTotal:
***************************查看ip
ip addr
em*
inet 192.168.1.116/24
***************************自动挂在硬盘
硬盘格式化sudo mkfs -t ext4 /dev/sda
https://blog.51cto.com/12348890/2092339
0.
mount /dev/sdb1 /mnt/newdisk
1.
df -hT
/dev/sdb1               xfs        15T  7.5T  7.2T   52% /newdisk
2.
/etc/fstab
/dev/sdb1 /newdisk    xfs    defaults    0 0
/dev/sdb1 /mnt/newdisk    xfs    defaults    0 0

***************************
添加用户： adduser djangouser
设置用户密码：passwd djangouser
更变一个目录或者文件的拥有者: chown liushuqing /home/data/
r-4 | w-2 | x-1
    
***************************读取/写入文件内容有中文，需要解码
,'r',encoding='UTF-8'
,'w',encoding='UTF-8'

***************************scp -r 传输文件加密码
yum install sshpass 没有先安装
安装完成，先scp -r -P 端口 源文件 目的文件密码调试 yes/no
sshpass -p 密码 scp -r -P 端口 源文件 目的文件 
小p密码；大P端口

***************************参数文件或者文件夹的绝对路径
parser = argparse.ArgumentParser(description = 'result2word')
parser.add_argument('--i',required=True, help = 'please input raw_result_file')
parser.add_argument('--lenMin', type = int, default=False,help = 'default=False') 默认是False，参数类型整型
parser.add_argument('--pn', type = int, default = 10,  help = 'please input fasta_parts_number;default = 10') 默认是10，参数类型整型
args = parser.parse_args()
path1=os.path.abspath(args.i) #文件的绝对路径

script_path = sys.path[0] #运行脚本的绝对路径

***************************.gz压缩文件内容写入.txt文件
barcode_path = 'barcode01_merged_01.fastq.gz'
fastq_file = 'barcode01.fastq'
with gzip.open(barcode_path, 'r') as f1, open(fastq_file, 'w') as f2:
    con = f1.readlines()
    for line in con:
         # 对从gz文件读取到的内容，应使用decode()进行转化，否则会报错：TypeError: write() argument must be str, not bytes
        f2.write(line.decode('utf-8'))
***************************cut以分隔符分割取第一行
cut -f1 -d'\t'

***************************pandas对某一列或者行进行处理
f = lambda x:int(x)
df["G3"]=df["G3"].map(f)

***************************属排序，对应属内的种二级排序
G_S_SN = AutoVivification()
G_GN = AutoVivification()
for i in bacteria_list:
    info_text = i.split('\t')
    G_S_SN[info_text[3]][info_text[7]] = int(info_text[9])
    G_GN[info_text[3]] = int(info_text[5])
S_queue = []  #最后排列的顺序
for i in sorted(G_GN.items(),key=operator.itemgetter(1), reverse=True):
    for j in sorted(G_S_SN[i[0]].items(), key=operator.itemgetter(1), reverse=True):
        S_queue.append(j[0])
***************************excel下拉菜单设置
https://baijiahao.baidu.com/s?id=1632493625731728693&wfr=spider&for=pc
    
***************************随机打乱列表
a = [1,5,7,9]
shuffle(a)

***************************window和linux路径/反义
result_path = r'H:\python3\project\test\1127\bfvp_new\data'
result_file = result_path+r'\bacteria_mock.txt'
window系统：result_file.split('\\')

***************************json处理
***写入json
import json
j = json.dumps(convert_list,ensure_ascii=False) #编码中文
with codecs.open('file.json', "w", "utf-8") as f:
    f.write(j)
***读取json
with open(result_file,'r',encoding='UTF-8') as fr:
    strJson = json.load(fr)
print(strJson['name'])

***************************特殊xls文件读取<少数xls文件因为本身文件格式问题，是不可以通过xlrd库读取>
f = open(result_file,'rb')
lines = f.readlines()
for line in lines:
    line = line.decode('gb2312') # 或者 line = line.decode('gb2312').encode('utf8')
    print(line)
    
***************************按次数count倒序，次数相同按id升序
id_list = [{'id': 123, 'count': 12, 'name': 'privacy'}, {'id': 47, 'count': 33, 'name': 'sensitive'},
           {'id': 109, 'count': 1, 'name': 'permission'}, {'id': 102, 'count': 3, 'name': 'sensitive'},
           {'id': 77, 'count': 3, 'name': 'sensitive'}, {'id': 52, 'count': 3, 'name': 'privacy'},
           {'id': 10, 'count': 1, 'name': 'permission'}]
print(sorted(id_list, key=lambda x:(-x['count'], int(x['id']))))

***************************程序选项（默认值）
parser.add_argument('--pn', default = 10, type = int, help = 'please input fasta_parts_number;default = 10')

***************************先读取第一行的标题栏（剔除掉标题栏）
fr.readline() 

***************************同时一次删除列表多个元素
list1 = [1,2,3,4,5,6,7,8,9,10]  #4,6
index_to_delete = [3,5]
list1 = [list1[i] for i in range(0, len(list1), 1) if i not in index_to_delete]

**************************深度复制
>>>import copy
>>> c = copy.deepcopy(a)

**************************dos2unix
断行符
window : ^M$ 
linux  : $
dos2unix filename

**************************split指定多个分隔符
import re
text='3.14:15'
print re.split('[.:]', text)

**************************软连接
ln -s 【目标目录】 【软链接地址】

**************************判断是文件还是文件夹？
import os
if os.path.isdir(path):
  print "it's a directory"
elif os.path.isfile(path):
  print "it's a normal file"
else:
  print "it's a special file(socket,FIFO,device file)"

**************************起始
str.startswith( 'this' )

**************************首行空格2字符
from docx import Document
from docx.shared import Cm

demo_file = r"t1.docx"
doc = Document(demo_file)
paragraph = doc.paragraphs[0]
paragraph.add_run('卷曲乳杆菌是革兰氏阳性杆菌，乳杆菌属。乳杆菌属是阴道和胃肠道自然菌群的主导成员，适量乳杆菌的存在可保持阴道的微环境，防止病原体侵入，使得微环境维持稳态，但是，过量的乳杆菌也是会引起疾病，如细胞溶解性阴道炎。乳杆菌菌株在发酵食品和饮料行业广泛使用，并用作疫苗输送系统，以激活动物模型的粘膜免疫等。一些研究表明，局部或口服某些乳杆菌菌株作为益生菌可预防阴道炎（复发性外阴道假丝酵母VVC）的复发，但仍缺乏其有效性的一致证据。乳杆菌属为革兰阳性杆菌，其大部分的病原菌是人和动物消化道黏膜表面共生菌群的一部分，主要存在于口腔、小肠、大肠、泌尿生殖道和皮肤表面。乳杆菌属是益生菌， ')
paragraph.paragraph_format.first_line_indent = Cm(0.74) #首行空格两个字符
#paragraph.paragraph_format.first_line_indent =406400
doc.save("{0}".format('test.docx'))

***************************提取文件大小
os.path.getsize('/data/liushuqing/test/Y2021/LC/0629_X101SC21022006-Z01-J021/seqData/spadesResult/TH1001/scaffolds.fasta')


 





