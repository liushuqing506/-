#!/usr/bin/python
# -*- coding: utf-8 -*-

import os
import time
import operator
from collections import OrderedDict
import argparse
import datetime
import sys
import re
import docx
import xlrd
import random
from docx import *
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx.oxml import OxmlElement
from PIL import Image,ImageFilter

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None
def delete_mark(doc_file,mark):
    paragraphs = list(doc_file.paragraphs)
    for i in range(len(paragraphs)):
        paragraph = paragraphs[i]
        if mark in paragraph.text:
            paragraph = doc.paragraphs[i]
            delete_paragraph(paragraph)

def add_headline(doc_file,mark,text1,text2,center=False,name='等线'):
    paragraphs = list(doc_file.paragraphs)
    for i in range(len(paragraphs)):
        paragraph = paragraphs[i]
        if mark in paragraph.text:
            paragraph = doc.paragraphs[i]
            paragraph.insert_paragraph_before('')
            paragraph = doc.paragraphs[i]
            run1 = paragraph.add_run(text1)
            run1.font.bold = True
            run1.font.color.rgb = RGBColor(18, 49, 58)
            run1.font.name = name
            s = run1._element
            s.rPr.rFonts.set(qn('w:eastAsia'), name)
            run1.font.italic = False
            run1.font.size = Pt(10)
            run2 = paragraph.add_run(text2)
            run2.font.bold = True
            run2.font.color.rgb = RGBColor(18, 49, 58)
            run2.font.name = name
            run2.font.italic = True
            run2.font.size = Pt(10)
            s = run2._element
            s.rPr.rFonts.set(qn('w:eastAsia'), name)
            paragraph.paragraph_format.line_spacing = 1.5
            if center == True:
                paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            else:
                pass


def add_annotate(doc_file,mark,text1,center=False,bold=False,name='等线'):
    paragraphs = list(doc_file.paragraphs)
    for i in range(len(paragraphs)):
        paragraph = paragraphs[i]
        if mark in paragraph.text:
            paragraph = doc.paragraphs[i]
            paragraph.insert_paragraph_before('')
            paragraph = doc.paragraphs[i]
            run1 = paragraph.add_run(text1)
            run1.font.bold = bold
            run1.font.color.rgb = RGBColor(0,0,0)
            run1.font.name = name
            s = run1._element
            s.rPr.rFonts.set(qn('w:eastAsia'), name)
            run1.font.italic = False
            run1.font.size = Pt(10)
            paragraph.paragraph_format.line_spacing = 1.5
            if center == True:
                paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            else:
                pass

def add_picture(doc_file,mark,picture):
    paragraphs = list(doc_file.paragraphs)
    for i in range(len(paragraphs)):
        paragraph = paragraphs[i]
        if mark in paragraph.text:
            paragraph = doc.paragraphs[i]
            paragraph.insert_paragraph_before('')
            paragraph = doc.paragraphs[i]
            paragraph.add_run().add_picture(picture)
            paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
