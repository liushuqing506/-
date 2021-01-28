import os
import time
import operator
from collections import OrderedDict
import argparse
import datetime
import sys
import re
import docx
import xlrd
import random
from docx import *
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx.oxml import OxmlElement
from PIL import Image,ImageFilter

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

def set_run(run, font_size=12, bold = False, R=0,G=0,B=0, name='等线',italic=False):
    '''
    设置run对象
    :param run:
    :param font_size: 字体大小
    :param bold: 是否加粗
    :param color: 字体颜色
    :param name: 字体名
    :return:
    '''
    run.font.size = Pt(font_size) #默认小四
    run.font.bold = bold  #默认不加粗
    run.font.color.rgb = RGBColor(R,G,B) #默认黑色
    run.font.name = name
    run.font.italic = italic #默认不斜体
    # 设置字体必须要下面2步
    s = run._element
    s.rPr.rFonts.set(qn('w:eastAsia'), name)
'''
def add_text(tables_num,row_loc,column_loc,text,italic=False,center=True,bold=False,R=0,G=0,B=0,name='等线',font_size=12):
    set_run(tables_num.cell(row_loc,column_loc).paragraphs[0].add_run(text),italic=italic,bold=bold,R=R,G=G,B=B,name=name,font_size=font_size)
    if center:
        tables_num.cell(row_loc,column_loc).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        tables_num.cell(row_loc,column_loc).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    else:
        pass

def table_frame(tables_num,rows_num):
    tables_num.add_row()
    for i in range(len(tables_num.columns)):
        set_cell_border(tables_num.cell(rows_num, i), \
                        top={"sz": 4, "val": "single", "color": RGBColor(49, 133, 156), }, \
                        bottom={"sz": 4, "color": RGBColor(49, 133, 156), "val": "single"}, \
                        left={"sz": 4, "val": "single", "color": RGBColor(49, 133, 156), }, \
                        right={"sz": 4, "color": RGBColor(49, 133, 156), "val": "single"}, \
                        insideH={"color": RGBColor(49, 133, 156)}, )


def table_add_row(tables_num,add_row_num):
    remove_row(tables_num, tables_num.rows[2])
#    add_row_num = 3  ##Change
    for i in range(add_row_num):
        table_frame(tables_num, 2 + i)  # add newrow and set frame
'''
def add_text(tables_num,row_loc,column_loc,text,italic=False,center=True,bold=False,R=0,G=0,B=0,name='等线',font_size=10):
    set_run(tables_num.cell(row_loc,column_loc).paragraphs[0].add_run(text),italic=italic,bold=bold,R=R,G=G,B=B,name=name,font_size=font_size)
    if center:
        tables_num.cell(row_loc,column_loc).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
#        tables_num.cell(row_loc,column_loc).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    else:
        pass
'''
def table_frame(tables_num,rows_num):
    tables_num.add_row()
    for i in range(len(tables_num.columns)):
        set_cell_border(tables_num.cell(rows_num, i), \
                        top={"sz": 4, "val": "single", "color": RGBColor(49, 133, 156), }, \
                        bottom={"sz": 4, "color": RGBColor(49, 133, 156), "val": "single"}, \
                        left={"sz": 4, "val": "single", "color": RGBColor(49, 133, 156), }, \
                        right={"sz": 4, "color": RGBColor(49, 133, 156), "val": "single"}, \
                        insideH={"color": RGBColor(49, 133, 156)}, )
'''
def table_frame(tables_num,rows_num,bottom=False):
    tables_num.add_row()
    if bottom:
        for i in range(len(tables_num.columns)):
            set_cell_border(tables_num.cell(rows_num, i), \
                bottom={"sz":12, "val": "single","color": RGBColor(3, 168, 158),},)
    else:
        pass

def table_add_row(tables_num,add_row_num):
    remove_row(tables_num, tables_num.rows[2])
    for i in range(add_row_num-1):
        table_frame(tables_num, 2 + i)  # add newrow and set frame
    table_frame(tables_num, 2+add_row_num-1,bottom=True)
